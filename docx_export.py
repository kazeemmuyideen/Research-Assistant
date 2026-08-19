import io
import re

from docx import Document
from docx.shared import Pt


def _add_structured_text(doc, text: str):
    """
    Parses the lightweight markdown the model is prompted to produce
    ('## Section Header' lines, '- bullet' / '1. numbered' lines, blank-line
    separated paragraphs) into real Word elements — actual Heading styles
    and List Bullet/List Number styles — instead of dumping raw text with
    literal '##' and '-' characters showing up on the page.
    """
    if not text:
        return

    paragraph_buffer = []

    def flush_paragraph():
        if paragraph_buffer:
            para_text = " ".join(paragraph_buffer).strip()
            if para_text:
                doc.add_paragraph(para_text)
            paragraph_buffer.clear()

    for raw_line in text.split("\n"):
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            continue

        header_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if header_match:
            flush_paragraph()
            level = min(len(header_match.group(1)) + 1, 4)  # offset since the doc title is level 0/1
            doc.add_heading(header_match.group(2).strip(), level=level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        if bullet_match:
            flush_paragraph()
            doc.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.*)", line)
        if numbered_match:
            flush_paragraph()
            doc.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            continue

        # Bold-line convention some models use for sub-headers, e.g. "**Key Facts**"
        bold_header_match = re.match(r"^\*\*(.+)\*\*$", line)
        if bold_header_match:
            flush_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(bold_header_match.group(1).strip())
            run.bold = True
            continue

        paragraph_buffer.append(line)

    flush_paragraph()


def build_research_docx(topic: str, summary: str, sources: list, tools_used: list, full_report: str = "") -> bytes:
    doc = Document()

    doc.add_heading(topic or "Research Report", level=0)

    doc.add_heading("Summary", level=1)
    _add_structured_text(doc, summary or "")

    if full_report and full_report.strip() != (summary or "").strip():
        doc.add_heading("Full Report", level=1)
        _add_structured_text(doc, full_report)

    if sources:
        doc.add_heading("Sources", level=1)
        for s in sources:
            doc.add_paragraph(s, style="List Bullet")

    if tools_used:
        doc.add_heading("Tools Used", level=1)
        doc.add_paragraph(", ".join(tools_used))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()